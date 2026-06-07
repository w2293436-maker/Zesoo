"""Word (DOCX) 导出服务"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(report: dict, output_dir: str) -> str:
    """根据报告 JSON 生成 Word 文档，返回文件路径"""
    doc = Document()

    # 默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(11)

    # ===== 封面 =====
    title = doc.add_heading("书籍精读报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    book_title = report.get("book_title", "未命名书籍")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"《{book_title}》")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(59, 130, 246)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(148, 163, 184)

    doc.add_paragraph()
    doc.add_paragraph("—" * 40)

    # ===== 逐章内容 =====
    chapters = report.get("chapters", [])
    for i, ch in enumerate(chapters, 1):
        chapter_name = ch.get("chapter_name", f"章节 {i}")

        # 章节标题
        doc.add_heading(f"{chapter_name}", level=1)

        # 1. 章节总结
        doc.add_heading("📝 章节总结", level=2)
        summary = ch.get("summary", "")
        if summary:
            sp = doc.add_paragraph(summary)
            sp.paragraph_format.space_after = Pt(6)

        # 2. 核心观点
        doc.add_heading("💡 核心观点", level=2)
        core_ideas = ch.get("core_ideas", [])
        if core_ideas:
            for j, item in enumerate(core_ideas, 1):
                idea_text = item.get("idea", "")
                elaboration = item.get("elaboration", "")
                p = doc.add_paragraph()
                run = p.add_run(f"观点 {j}：")
                run.bold = True
                p.add_run(idea_text)
                if elaboration:
                    p2 = doc.add_paragraph(elaboration)
                    p2.paragraph_format.left_indent = Cm(0.5)
                    for r in p2.runs:
                        r.font.color.rgb = RGBColor(75, 85, 99)
                        r.font.size = Pt(10)
        else:
            doc.add_paragraph("（无）")

        # 3. 金句摘录
        doc.add_heading("✨ 金句摘录", level=2)
        quotes = ch.get("quotes", [])
        if quotes:
            for q in quotes:
                text = q.get("text", "")
                context = q.get("context", "")
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(f'"{text}"')
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(55, 65, 81)
                if context:
                    ctx_para = doc.add_paragraph()
                    ctx_para.paragraph_format.left_indent = Cm(0.5)
                    run = ctx_para.add_run(f"—— {context}")
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(148, 163, 184)
                doc.add_paragraph()
        else:
            doc.add_paragraph("（无）")

        # 4. 方法论
        doc.add_heading("🔧 方法论", level=2)
        methodology = ch.get("methodology", [])
        if methodology:
            for m in methodology:
                name = m.get("name", "")
                desc = m.get("description", "")
                steps = m.get("steps", [])
                doc.add_heading(name, level=3)
                if desc:
                    p_desc = doc.add_paragraph(desc)
                    for r in p_desc.runs:
                        r.font.size = Pt(10)
                if steps:
                    p = doc.add_paragraph()
                    run = p.add_run("操作步骤：")
                    run.bold = True
                    run.font.size = Pt(10)
                    for k, step in enumerate(steps, 1):
                        sp = doc.add_paragraph(f"{k}. {step}")
                        sp.paragraph_format.left_indent = Cm(0.5)
                        for r in sp.runs:
                            r.font.size = Pt(10)
        else:
            doc.add_paragraph("（无）")

        # 5. 启示建议
        doc.add_heading("🎯 启示与行动建议", level=2)
        insights = ch.get("actionable_insights", [])
        if insights:
            for item in insights:
                insight_text = item.get("insight", "")
                action_text = item.get("action", "")
                p = doc.add_paragraph()
                run = p.add_run("💡 ")
                p.add_run(insight_text)
                if action_text:
                    p2 = doc.add_paragraph()
                    run = p2.add_run("行动建议：")
                    run.bold = True
                    run.font.size = Pt(10)
                    run2 = p2.add_run(action_text)
                    run2.font.size = Pt(10)
                    p2.paragraph_format.left_indent = Cm(0.5)
                doc.add_paragraph()
        else:
            doc.add_paragraph("（无）")

        # 章间分隔（最后一章不加）
        if i < len(chapters):
            doc.add_page_break()

    # ===== 免责声明页 =====
    doc.add_page_break()
    doc.add_heading("免责声明", level=1)

    disclaimers = [
        ("AI 生成内容声明",
         "本报告由择书Zesoo基于人工智能技术自动生成，所有内容（包括章节总结、核心观点、"
         "金句摘录、方法论、启示建议等）仅供参考和学习交流之用，不构成任何形式的专业建议、"
         "指导或承诺。AI 生成的内容可能存在不准确、不完整或与原文存在偏差的情况，"
         "用户应自行判断和核实报告内容的准确性和适用性。"),
        ("知识产权声明",
         "用户上传的书籍文件及其内容的知识产权归原作者及权利人所有。本工具仅提供文本分析"
         "和内容提炼服务，不对上传文件的内容主张任何权利。本报告仅供用户个人学习使用，"
         "用户不得将本报告用于商业目的，亦不得以任何形式公开发布或传播涉及侵权的内容。"),
        ("责任限制",
         "本工具按「现状」提供，不提供任何明示或默示的保证。在任何情况下，本工具的开发者和"
         "运营方均不对因使用或无法使用本工具所产生的任何直接、间接、附带、特殊或后果性"
         "损失承担责任，包括但不限于数据丢失、业务中断、信息不准确导致的决策失误等。"
         "使用本工具的风险由用户自行承担。"),
    ]

    for title, content in disclaimers:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph(content)
        p.paragraph_format.space_after = Pt(12)

    # ===== 页脚 =====
    doc.add_paragraph()
    doc.add_paragraph("—" * 40)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("本报告由 AI 自动生成，仅供参考 | 择书Zesoo")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(203, 213, 225)

    # 保存
    filename = f"书籍精读报告_{book_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    return filepath
